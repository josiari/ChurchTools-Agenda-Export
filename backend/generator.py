import dateutil.parser
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    """Set the background color of a cell."""

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_cell_run(cell, text, color=RGBColor(0, 0, 0)):
    """Add a run to a cell with specified text and color."""

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = 'Montserrat Light'
    run.font.size = Pt(12)
    run.font.color.rgb = color
    return run


def get_name_of_service(event, service_id):
    """Get the name of a service by its ID."""

    return ', '.join(
        service['name']
        for service in event['data']['eventServices']
        if service['serviceId'] == service_id and service['name']
    )


def add_info_table(doc, event):
    """Add a table with general information to the document."""

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in table.rows:
        row.height = Cm(0.7)
        for cell in row.cells:
            set_cell_background(cell, 'EEEEEE')
            cell.width = Cm(8)

    info = [
        ("Leitung", 3),
        ("Predigt", 1),
        ("Band", 2),
        ("Ton", 6),
        ("Video", 32),
        ("Folien", 7),
    ]

    for i, (label, service_id) in enumerate(info[:3]):
        add_cell_run(table.cell(i, 0), f"{label}: {get_name_of_service(event, service_id)}")
    for i, (label, service_id) in enumerate(info[3:]):
        add_cell_run(table.cell(i, 1), f"{label}: {get_name_of_service(event, service_id)}")

    add_cell_run(table.cell(3, 0).merge(table.cell(3, 1)), "Notizen: \n\n")

def add_agenda_table(doc, agenda):
    """Add an agenda table to the document."""

    # Create a table for the agenda
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths and header
    col_width = [1.9, 4.5, 2.5, 8.53]
    col_header = ['Uhrzeit', 'Programmpunkt', 'Zuständig', 'Notizen']

    # Add header row
    for i, cell in enumerate(table.rows[0].cells):
        set_cell_background(cell, '666666')
        add_cell_run(cell, col_header[i], RGBColor(255, 255, 255))

    # Add agenda items
    for item in agenda['data']['items']:
        row = table.add_row()
        if item['type'] == 'header':
            merged_cell = row.cells[0].merge(row.cells[1].merge(row.cells[2].merge(row.cells[3])))
            set_cell_background(merged_cell, 'EEEEEE')
            add_cell_run(merged_cell, item['title'])
        else:
            values = [
                dateutil.parser.parse(item['start']).strftime('%H:%M') if 'start' in item else '',
                item['title'],
                '[Band]' if item['type'] == 'song' else item['responsible']['text'],
                item.get('note', ''),
            ]
            for idx, val in enumerate(values):
                add_cell_run(row.cells[idx], val)

    # Adjust row heights and column widths
    for row in table.rows:
        row.height = Cm(0.7)
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Cm(col_width[idx])

def generate_docx(event, agenda):
    """Generate a Word document for the event agenda."""

    doc = Document()

    # Set margins
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Add title
    title = doc.add_heading(
        f"Gottesdienst Agenda {dateutil.parser.parse(event['data']['startDate']).strftime('%d.%m.%Y')}",
        level=0
    )
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # Add space

    # Add general information table
    add_info_table(doc, event)

    doc.add_paragraph()  # Add space

    # Add agenda table
    add_agenda_table(doc, agenda)

    return doc